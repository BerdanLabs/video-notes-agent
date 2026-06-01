from pathlib import Path

from docx import Document

from video_notes.qa import check_notes


def test_check_docx_flags_php_artifact(tmp_path: Path):
    path = tmp_path / "bad.docx"
    doc = Document()
    doc.add_paragraph("Why does this end with PHP")
    doc.save(path)

    report = check_notes(path)

    assert report["paragraphs"] == 1
    assert report["problems"]
    assert "Suspicious PHP artifact" in report["problems"][0]


def test_check_docx_accepts_normal_questions(tmp_path: Path):
    path = tmp_path / "good.docx"
    doc = Document()
    doc.add_paragraph("Why does copywriting matter?")
    doc.save(path)

    report = check_notes(path)

    assert report["problems"] == []


def test_check_markdown_flags_php_artifact(tmp_path: Path):
    path = tmp_path / "bad.md"
    path.write_text("This line contains a PHP artifact", encoding="utf-8")

    report = check_notes(path)

    assert report["paragraphs"] == 1
    assert report["problems"]
    assert "Suspicious PHP artifact" in report["problems"][0]


def test_check_markdown_flags_bad_patterns(tmp_path: Path):
    path = tmp_path / "corrupted.md"
    path.write_text("# Notes\n\nThis contains ?why smart quote issues\n\n![Screenshot](shot.png)", encoding="utf-8")

    report = check_notes(path)

    assert report["headings"] == 1
    assert report["paragraphs"] == 1
    assert report["images"] == 1
    assert len(report["problems"]) == 1
    assert "Found suspicious pattern: ?why" in report["problems"][0]
