from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .synthesis import build_note_plan
from .utils import clean_docx_text, fmt_time


def build_docx(
    title: str,
    source: str,
    transcript: list[dict],
    screenshots: list[Path],
    out_path: Path,
    profile: str = "course",
) -> Path:
    doc = Document()
    setup_styles(doc)

    doc.add_paragraph(clean_docx_text(title), style="Title")
    meta = doc.add_paragraph()
    meta.add_run("Source: ").bold = True
    meta.add_run(clean_docx_text(source))

    note_plan = build_note_plan(transcript, profile)

    doc.add_heading("Quick Summary", level=1)
    for bullet in note_plan["summary"]:
        doc.add_paragraph(clean_docx_text(bullet), style="List Bullet")

    doc.add_heading("Memorable Lines", level=1)
    for segment in note_plan["quotes"]:
        p = doc.add_paragraph(style="Timestamp Quote")
        p.add_run(f'{fmt_time(segment["start"])} - "{clean_docx_text(segment["text"])}"')

    doc.add_heading("Detailed Notes", level=1)
    if screenshots:
        add_screenshot(doc, screenshots[0], "Opening visual from the video.")
    for heading, bullets in note_plan["sections"].items():
        doc.add_heading(heading, level=2)
        for bullet in bullets:
            doc.add_paragraph(clean_docx_text(bullet), style="List Bullet")

    if screenshots[1:]:
        doc.add_heading("Selected Screenshots", level=1)
        for shot in screenshots[1:]:
            add_screenshot(doc, shot, f"Selected frame: {shot.stem.replace('-', ':')}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 23, "6F1026"),
        ("Heading 1", 16, "6F1026"),
        ("Heading 2", 13, "222222"),
    ]:
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    if "Timestamp Quote" not in styles:
        quote = styles.add_style("Timestamp Quote", WD_STYLE_TYPE.PARAGRAPH)
        quote.font.name = "Aptos"
        quote.font.size = Pt(9.5)
        quote.font.italic = True
        quote.font.color.rgb = RGBColor(65, 65, 65)


def add_screenshot(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.25))
    cap = doc.add_paragraph(clean_docx_text(caption))
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def summarize_placeholder(transcript: list[dict], profile: str) -> list[str]:
    return build_note_plan(transcript, profile)["summary"]
