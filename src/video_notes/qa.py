from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


BAD_PATTERNS = ["?s", "?why", "?poets?", "?killers?", "\ufffd"]
INTERNAL_PATTERNS = ["What Improved", "Comparison With", "Coverage note", "QA caveat"]


def check_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_text = [cell.text for t in doc.tables for row in t.rows for cell in row.cells]
    text = "\n".join(paragraphs + table_text)
    with ZipFile(path) as zf:
        images = [n for n in zf.namelist() if n.startswith("word/media/")]
    problems = []
    for pattern in BAD_PATTERNS + INTERNAL_PATTERNS:
        if pattern in text:
            problems.append(f"Found suspicious pattern: {pattern}")
    for line in text.splitlines():
        if "PHP" in line and not any(ch.isdigit() for ch in line):
            problems.append(f"Suspicious PHP artifact: {line[:120]}")
    return {
        "paragraphs": len(paragraphs),
        "headings": len([p for p in doc.paragraphs if p.style.name.startswith("Heading")]),
        "tables": len(doc.tables),
        "images": len(images),
        "problems": problems,
    }
